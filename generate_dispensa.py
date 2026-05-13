"""Genera la dispensa unica del corso 16h come docx custom.

Caratteristiche:
- Cover page con banner grafico
- TOC automatico
- Heading colorati H1/H2/H3
- Callout box (info, warning, tip, danger, success)
- Code block stilizzati
- Tabelle con header colorato
- Immagini (terminale, code, diagrammi) generate da img/
- Spaziatura aumentata
"""

import os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement


# === COLORI ===
COL_PRIMARY = "0F2D52"       # blu scuro
COL_ACCENT = "00A0B0"        # teal
COL_DANGER = "E63946"        # rosso
COL_OK = "06A77D"            # verde
COL_WARN = "F7931E"          # arancio
COL_INFO = "3B82F6"          # blu chiaro info
COL_TEXT = "1A1A2E"
COL_GREY = "6B7280"
COL_LIGHT_GREY = "F1F4F8"


def hex_to_rgb(hex_str):
    return RGBColor(int(hex_str[0:2], 16), int(hex_str[2:4], 16),
                     int(hex_str[4:6], 16))


def set_cell_bg(cell, hex_color):
    """Sfondo cella tabella."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_cell_borders(cell, color="CCCCCC", size="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for border in ["top", "left", "bottom", "right"]:
        b = OxmlElement(f"w:{border}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), size)
        b.set(qn("w:color"), color)
        tc_borders.append(b)
    tc_pr.append(tc_borders)


# === SETUP STYLES ===
def setup_styles(doc):
    """Personalizza stili Heading/Normal."""
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = hex_to_rgb(COL_TEXT)
    pf = style.paragraph_format
    pf.space_after = Pt(8)
    pf.line_spacing = 1.5

    # H1
    h1 = doc.styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(28)
    h1.font.bold = True
    h1.font.color.rgb = hex_to_rgb(COL_PRIMARY)
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(12)
    h1.paragraph_format.line_spacing = 1.2

    # H2
    h2 = doc.styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(20)
    h2.font.bold = True
    h2.font.color.rgb = hex_to_rgb(COL_ACCENT)
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(8)
    h2.paragraph_format.line_spacing = 1.2

    # H3
    h3 = doc.styles["Heading 3"]
    h3.font.name = "Calibri"
    h3.font.size = Pt(15)
    h3.font.bold = True
    h3.font.color.rgb = hex_to_rgb(COL_PRIMARY)
    h3.paragraph_format.space_before = Pt(14)
    h3.paragraph_format.space_after = Pt(6)


def set_margins(doc, top=2.5, bottom=2.5, left=2.5, right=2.5):
    for section in doc.sections:
        section.top_margin = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin = Cm(left)
        section.right_margin = Cm(right)


# === ELEMENTI ===
def add_cover(doc):
    """Pagina di copertina con banner immagine."""
    section = doc.sections[0]
    # Cover: senza margini grandi
    sect_pr = section._sectPr
    # Riduci temporaneamente i margini per la cover
    section.top_margin = Cm(0)
    section.left_margin = Cm(0)
    section.right_margin = Cm(0)

    # Banner
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture("img/cover_banner.png", width=Cm(21.0))

    # Spazio
    for _ in range(2):
        doc.add_paragraph()

    # Titolo
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Modulo Secure Coding")
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = hex_to_rgb(COL_PRIMARY)
    run.font.name = "Calibri"

    # Sottotitolo
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Dispensa completa · 16 ore")
    run.font.size = Pt(22)
    run.font.color.rgb = hex_to_rgb(COL_ACCENT)
    run.font.italic = True
    run.font.name = "Calibri"

    # Spazio
    for _ in range(3):
        doc.add_paragraph()

    # Linea decorativa
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("━" * 30)
    run.font.color.rgb = hex_to_rgb(COL_ACCENT)
    run.font.size = Pt(14)

    # Box informazioni
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Corso IFTS STEM")
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = hex_to_rgb(COL_PRIMARY)
    run.font.name = "Calibri"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Specialista nelle Tecniche di Evoluzione e Manutenzione del Software")
    run.font.size = Pt(12)
    run.font.italic = True
    run.font.color.rgb = hex_to_rgb(COL_GREY)
    run.font.name = "Calibri"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\nCodice Progetto 321386 · Matricola 2025IS0766")
    run.font.size = Pt(11)
    run.font.color.rgb = hex_to_rgb(COL_GREY)
    run.font.name = "Calibri"

    # Spazio prima del footer
    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Ing. Alessandro Manneschi")
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = hex_to_rgb(COL_PRIMARY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Anno formativo 2024/2025 · Versione 2.0")
    run.font.size = Pt(11)
    run.font.color.rgb = hex_to_rgb(COL_GREY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\nAssoservizi · ITS Prodigi · Polo Tecnologico Manetti Porciatti · Università di Siena · Opus Automazione")
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = hex_to_rgb(COL_GREY)

    # Page break
    doc.add_page_break()

    # Ripristino margini per il resto del documento
    section.top_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


def add_toc(doc):
    """Aggiunge un TOC automatico (Word aggiorna a F9)."""
    h = doc.add_paragraph()
    run = h.add_run("Indice")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = hex_to_rgb(COL_PRIMARY)
    h.paragraph_format.space_after = Pt(18)

    p = doc.add_paragraph()
    run = p.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")

    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'

    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")

    fldChar3 = OxmlElement("w:t")
    fldChar3.text = "[Premi F9 in Word per aggiornare l'indice]"

    fldChar4 = OxmlElement("w:fldChar")
    fldChar4.set(qn("w:fldCharType"), "end")

    r_element = run._r
    r_element.append(fldChar1)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar3)
    r_element.append(fldChar4)

    doc.add_page_break()


def add_chapter_cover(doc, num, title, subtitle, hours):
    """Apertura di capitolo con grafica."""
    # Banda colorata simulata con un paragrafo riempito
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(60)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(f"CAPITOLO {num}")
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = hex_to_rgb(COL_ACCENT)
    run.font.name = "Calibri"

    # Linea decorativa
    p = doc.add_paragraph()
    run = p.add_run("━" * 8)
    run.font.color.rgb = hex_to_rgb(COL_ACCENT)
    run.font.size = Pt(14)
    p.paragraph_format.space_after = Pt(24)

    # Titolo
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = hex_to_rgb(COL_PRIMARY)
    run.font.name = "Calibri"
    p.paragraph_format.space_after = Pt(12)

    # Sottotitolo
    p = doc.add_paragraph()
    run = p.add_run(subtitle)
    run.font.size = Pt(16)
    run.font.italic = True
    run.font.color.rgb = hex_to_rgb(COL_GREY)
    p.paragraph_format.space_after = Pt(24)

    # Durata
    p = doc.add_paragraph()
    run = p.add_run(f"⏱  {hours}")
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = hex_to_rgb(COL_ACCENT)

    doc.add_page_break()


def add_h1(doc, text):
    doc.add_heading(text, level=1)


def add_h2(doc, text):
    doc.add_heading(text, level=2)


def add_h3(doc, text):
    doc.add_heading(text, level=3)


def add_para(doc, text, italic=False, bold=False, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.italic = italic
    run.font.bold = bold
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(1.0)
    run = p.runs[0] if p.runs else p.add_run()
    run.text = text
    run.font.size = Pt(11)
    return p


def add_image(doc, path, caption=None, width_cm=15):
    if not os.path.exists(path):
        print(f"  Missing: {path}")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run()
    run.add_picture(path, width=Cm(width_cm))

    if caption:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Fig. — {caption}")
        run.font.size = Pt(10)
        run.font.italic = True
        run.font.color.rgb = hex_to_rgb(COL_GREY)
        p.paragraph_format.space_after = Pt(16)


def add_callout(doc, kind: str, text: str):
    """Box callout colorato (info, warning, tip, danger, success)."""
    color_map = {
        "info": (COL_INFO, "ℹ", "Info"),
        "warning": (COL_WARN, "⚠", "Attenzione"),
        "tip": (COL_OK, "💡", "Suggerimento"),
        "danger": (COL_DANGER, "🚩", "Pericolo"),
        "success": (COL_OK, "✓", "Best practice"),
        "story": (COL_PRIMARY, "📖", "Una storia"),
    }
    color, icon, label = color_map.get(kind, color_map["info"])

    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, "FAFBFC")
    set_cell_borders(cell, color=color, size="12")

    # Padding interno
    cell_p = cell.paragraphs[0]
    cell_p.paragraph_format.space_before = Pt(4)

    # Label
    p_lbl = cell.paragraphs[0]
    run = p_lbl.add_run(f"{icon}  {label}")
    run.font.bold = True
    run.font.color.rgb = hex_to_rgb(color)
    run.font.size = Pt(11)

    # Testo
    p = cell.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = hex_to_rgb(COL_TEXT)

    # Spaziatura sopra/sotto la tabella
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_code_block(doc, code: str, lang: str = "python"):
    """Blocco di codice con sfondo grigio chiaro e monospaziato."""
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, "1E1E28")    # sfondo scuro
    set_cell_borders(cell, color=COL_ACCENT, size="8")

    # Rimuovi paragrafo default
    cell.paragraphs[0].text = ""

    for line in code.split("\n"):
        p = cell.add_paragraph() if line != code.split("\n")[0] else cell.paragraphs[0]
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line if line else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(10)
        run.font.color.rgb = hex_to_rgb("E5E5E5")

    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_styled_table(doc, headers, rows, col_widths=None):
    """Tabella con header colorato + righe alternate."""
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    hdr_cells = tbl.rows[0].cells
    for i, h in enumerate(headers):
        cell = hdr_cells[i]
        set_cell_bg(cell, COL_PRIMARY)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = cell.paragraphs[0].add_run(h)
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = hex_to_rgb("FFFFFF")
        run.font.name = "Calibri"

    # Rows
    for r_idx, row in enumerate(rows):
        row_cells = tbl.rows[r_idx + 1].cells
        bg = "F8F9FB" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row):
            cell = row_cells[c_idx]
            set_cell_bg(cell, bg)
            set_cell_borders(cell, color="E0E0E0", size="4")
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(10)
            run.font.color.rgb = hex_to_rgb(COL_TEXT)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)


# ============================================================================
# CONTENUTO DISPENSA
# ============================================================================

def build_dispensa():
    doc = Document()
    setup_styles(doc)
    set_margins(doc)

    # === COVER ===
    add_cover(doc)

    # === TOC ===
    add_toc(doc)

    # === INTRODUZIONE ===
    add_h1(doc, "Introduzione")
    add_para(doc,
        "Stai per leggere una dispensa di secure coding: come si scrive codice "
        "che non venga bucato. Non è un manuale teorico astratto, e non è un "
        "libro di pentest 'rompiamo tutto e basta'. È una via di mezzo "
        "operativa: ti faccio vedere come gli attaccanti pensano, ma soprattutto "
        "ti insegno a scrivere codice che resista ai loro attacchi.")

    add_h2(doc, "A chi è rivolta")
    add_para(doc,
        "A te che stai facendo l'IFTS STEM. Hai già visto Python, ti stai "
        "affacciando a Java, JavaScript, PHP. Conosci la sintassi, sai scrivere "
        "uno script che gira, magari un mini web service con Flask. Quello che "
        "probabilmente non sai ancora è che lo stesso codice 'funzionante' può "
        "essere un disastro di sicurezza.")

    add_h2(doc, "Come è organizzata")
    add_para(doc,
        "8 capitoli più appendici, uno per ogni lezione di 2 ore. Ogni "
        "capitolo segue lo stesso pattern:")
    for b in [
        "Cosa imparerai — gli obiettivi della lezione in 3-5 punti.",
        "Una storia per cominciare — un caso reale o un'analogia.",
        "Spiegazione passo passo — i concetti tecnici con esempi.",
        "Laboratorio — cosa fai in aula (e puoi rifare a casa).",
        "Cosa portarti via — i 3-5 takeaway memorabili.",
        "Errori comuni — quelli che fanno tutti i junior, da non fare tu.",
    ]:
        add_bullet(doc, b)

    add_callout(doc, "info",
        "La sicurezza non è un add-on. Non è 'una cosa che il senior aggiunge "
        "alla fine'. È un modo di scrivere codice. Se la impari ora, ti porterai "
        "un vantaggio enorme rispetto a chi la imparerà dopo il primo breach.")

    doc.add_page_break()

    # === CAPITOLO 0 — SETUP ===
    add_chapter_cover(doc, "0", "Preparare l'ambiente di lavoro",
                       "Python, VS Code, terminale, virtual environment, Flask",
                       "30 minuti · da fare a casa o all'inizio della prima lezione")

    add_para(doc,
        "Prima di parlare di sicurezza, devi avere un ambiente che funziona. "
        "Questo capitolo serve a quello. Non è 'sicurezza', ma se salti queste "
        "30 minuti, le 16 ore successive saranno frustranti.")

    add_h2(doc, "Cosa installeremo")
    add_para(doc, "Cinque strumenti, tutti gratuiti, tutti multipiattaforma:")
    for b in [
        "Python 3.12 — il linguaggio principale dei nostri laboratori.",
        "Visual Studio Code — l'editor di codice, gratuito di Microsoft.",
        "Git — per scaricare il codice di esempio dal repository del corso.",
        "DB Browser for SQLite — un'interfaccia grafica per 'sbirciare' nei database.",
        "Browser moderno — probabilmente lo hai già.",
    ]:
        add_bullet(doc, b)

    add_h2(doc, "Installare Python")

    add_h3(doc, "Su Windows")
    for b in [
        "Vai su python.org/downloads",
        "Clicca sul pulsante giallo 'Download Python 3.12.x'",
        "Apri il file scaricato (python-3.12.x-amd64.exe)",
        "Importantissimo: spunta 'Add Python to PATH' nella prima schermata",
        "Clicca 'Install Now' e aspetta che finisca",
    ]:
        add_bullet(doc, b)

    add_callout(doc, "warning",
        "Se dimentichi di spuntare 'Add Python to PATH' durante l'installazione, "
        "nulla funzionerà dopo. Se ti capita: disinstalla Python e reinstallalo "
        "ricordandoti la casella.")

    add_h3(doc, "Verifica che funzioni")
    add_para(doc,
        "Apri un nuovo terminale (PowerShell su Windows, Terminal su macOS, "
        "qualsiasi shell su Linux) e digita:")
    add_image(doc, "img/cap0_python_version.png",
              caption="Verifica installazione Python e creazione venv",
              width_cm=14)

    add_h2(doc, "Installare VS Code e prima app Flask")
    add_para(doc,
        "Scarica VS Code da code.visualstudio.com, installa l'estensione Python. "
        "Quindi crea una cartella per il corso e installa Flask:")
    add_image(doc, "img/cap0_flask_run.png",
              caption="Installazione Flask e avvio del primo server web",
              width_cm=14)

    add_callout(doc, "tip",
        "Routine quotidiana: dopo il primo setup, ogni volta che apri il "
        "computer per il corso basta: 1) terminale, 2) cd cartella-progetto, "
        "3) attiva venv, 4) code . per aprire VS Code. Tre comandi sempre uguali.")

    doc.add_page_break()

    # === CAPITOLO 1 — PERCHÉ SECURE CODING ===
    add_chapter_cover(doc, "1", "Perché il secure coding",
                       "Fondamenti, CIA Triad, 5 principi, mentalità avversaria",
                       "Lezione 1 · 2 ore")

    add_h2(doc, "Cosa imparerai")
    for b in [
        "Cosa significa 'sicurezza' in informatica e come si misura.",
        "La differenza tra sicurezza del codice e sicurezza infrastrutturale.",
        "I cinque principi fondamentali del secure coding.",
        "Tre casi reali di breach causati da errori di codice.",
        "Cos'è la 'mentalità avversaria' e perché ti serve.",
    ]:
        add_bullet(doc, b)

    add_h2(doc, "Una storia per cominciare: Equifax 2017")

    add_callout(doc, "story",
        "Nel marzo 2017 viene pubblicata una correzione (patch) per Apache "
        "Struts, una libreria che molti siti web usano. La vulnerabilità che "
        "la patch corregge si chiama CVE-2017-5638 ed è valutata 10 su 10 di "
        "gravità: chi non installa la patch può essere completamente "
        "compromesso da chiunque su Internet.")

    add_para(doc,
        "Equifax è un'agenzia di credito americana, una delle più grandi al "
        "mondo. Tratta dati di centinaia di milioni di persone: nome, cognome, "
        "codice fiscale (Social Security Number), data di nascita, indirizzo. "
        "Equifax usa Apache Struts. Equifax è obbligata, per contratto e per "
        "legge, a installare le patch di sicurezza in tempi rapidi.")

    add_para(doc, "Equifax non installa la patch per due mesi.", bold=True)

    add_para(doc,
        "Tra il maggio e il luglio 2017, attaccanti sfruttano la vulnerabilità "
        "e rubano i dati di 147 milioni di persone. La metà degli adulti "
        "americani. Il costo finale per Equifax? Circa 1,4 miliardi di dollari "
        "tra multe, class action, settlement, perdita di valore in borsa, "
        "licenziamenti dei dirigenti.")

    add_callout(doc, "tip",
        "La sicurezza non si fa il giorno del breach. Si fa nei mesi e negli "
        "anni prima. È un'abitudine quotidiana di chi scrive codice.")

    add_h2(doc, "Cosa significa 'essere sicuri'")
    add_para(doc,
        "In informatica, 'sicurezza' non è una cosa sola. È la conservazione di "
        "proprietà misurabili. Le tre più importanti si chiamano CIA Triad:")

    add_image(doc, "img/cap1_cia_triad.png",
              caption="Le 3 proprietà fondamentali della sicurezza",
              width_cm=15)

    add_para(doc,
        "Ogni breach viola almeno una di queste tre proprietà. Equifax ha "
        "violato la C (dati rubati). Un attacco DDoS che mette KO un sito viola "
        "la A. Un bonifico modificato in transito viola la I.")

    add_h2(doc, "Sicurezza del codice vs sicurezza informatica")

    add_styled_table(doc,
        ["", "Sicurezza informatica", "Sicurezza del codice"],
        [
            ["Cosa protegge", "Rete, server, OS", "Applicazioni, API, dati"],
            ["Chi previene", "Sistemisti, IT", "SVILUPPATORI (tu)"],
            ["Esempi di difesa", "Firewall, VPN, MFA su SSH", "Validation, query parametrizzate, escape"],
            ["Esempi di attacco", "SSH brute force, port scan", "SQL Injection, XSS, IDOR"],
        ])

    add_callout(doc, "warning",
        "La frase più pericolosa che senti in azienda è: 'abbiamo il firewall, "
        "siamo a posto'. Falsa: una SQL Injection passa attraverso la porta 443 "
        "che il firewall lascia aperta apposta. Allo stesso modo: 'abbiamo "
        "HTTPS' protegge il CANALE, non l'applicazione dentro.")

    add_h2(doc, "I cinque principi del secure coding")
    add_para(doc,
        "Ci sono cinque principi che, se rispetti, eviti la stragrande maggioranza "
        "dei guai. Memorizzali, scrivili in un post-it:")

    add_styled_table(doc,
        ["#", "Principio", "In sintesi"],
        [
            ["1", "Least Privilege", "Ogni componente ha il MINIMO privilegio necessario"],
            ["2", "Defense in Depth", "Mai una difesa sola: più strati indipendenti"],
            ["3", "Fail Secure", "Quando qualcosa va storto, il sistema chiude (default deny)"],
            ["4", "KISS", "Meno codice = meno bug = meno superficie d'attacco"],
            ["5", "Separation of Duties", "Nessun singolo componente da solo per azioni critiche"],
        ])

    add_image(doc, "img/cap1_defense_depth.png",
              caption="Defense in Depth: strati di difesa indipendenti",
              width_cm=14)

    add_h2(doc, "La mentalità avversaria")
    add_para(doc,
        "Quando guardi del codice come sviluppatore, ti chiedi: 'fa quello che "
        "deve fare?'. Quando guardi del codice come attaccante, ti chiedi: "
        "'cosa fa che NON dovrebbe fare?'. Sembra la stessa cosa: non lo è.")

    add_callout(doc, "tip",
        "La mentalità avversaria non è essere paranoici. È solo l'altra metà "
        "della professionalità. Ogni riga di codice che scrivi, chiediti: cosa "
        "fa che non dovrebbe fare?")

    doc.add_page_break()

    # === CAPITOLO 2 — OWASP + STRIDE ===
    add_chapter_cover(doc, "2", "OWASP, threat modeling e STRIDE",
                       "Pensare alla sicurezza prima di scrivere codice",
                       "Lezione 2 · 2 ore")

    add_h2(doc, "Cosa imparerai")
    for b in [
        "Cos'è OWASP e perché tutti la citano.",
        "Le dieci vulnerabilità più diffuse nel web (OWASP Top 10).",
        "Cos'è una CVE e come si legge un punteggio CVSS.",
        "Come si fa un threat modeling leggero in 30 minuti.",
        "Cos'è STRIDE e come applicarlo con un esempio concreto.",
    ]:
        add_bullet(doc, b)

    add_h2(doc, "OWASP: chi sono")
    add_para(doc,
        "OWASP sta per Open Web Application Security Project. È una fondazione "
        "no-profit, internazionale, fondata nel 2001. Non vende nulla. "
        "Pubblica documenti gratuiti, tool gratuiti, lezioni gratuite. "
        "È LO STANDARD della sicurezza applicativa nel mondo.")

    add_h2(doc, "OWASP Top 10 (2021/2025)")
    add_styled_table(doc,
        ["#", "Vulnerabilità", "Esempio"],
        [
            ["A01", "Broken Access Control", "Cambi ?id=42 in ?id=43 e vedi dati altrui"],
            ["A02", "Cryptographic Failures", "Password in MD5, HTTPS mancante"],
            ["A03", "Injection", "SQL Injection, XSS"],
            ["A04", "Insecure Design", "Manca threat modeling"],
            ["A05", "Security Misconfiguration", "debug=True in produzione"],
            ["A06", "Vulnerable Components", "Libreria con CVE nota (Log4Shell)"],
            ["A07", "Auth Failures", "Login senza rate limit, no MFA"],
            ["A08", "Data Integrity", "Aggiornamenti non verificati"],
            ["A09", "Logging Failures", "Non rilevi il breach per 200 giorni"],
            ["A10", "SSRF", "App fetcha URL fornito dall'utente"],
        ])

    add_h2(doc, "CVE e CVSS in due minuti")
    add_para(doc,
        "CVE sta per Common Vulnerabilities and Exposures. È l'identificatore "
        "univoco di una vulnerabilità nota. Formato: CVE-AAAA-NNNNN (anno + "
        "progressivo).")
    add_para(doc, "Esempi famosi:")
    for b in [
        "CVE-2014-0160 → Heartbleed",
        "CVE-2017-5638 → Apache Struts (Equifax)",
        "CVE-2021-44228 → Log4Shell",
        "CVE-2024-3094 → XZ Utils backdoor",
    ]:
        add_bullet(doc, b)

    add_para(doc, "Il punteggio CVSS misura la gravità da 0 a 10:")
    add_styled_table(doc,
        ["Punteggio", "Severity", "Azione tipica"],
        [
            ["0.0", "None", "—"],
            ["0.1 – 3.9", "Low", "Programma il fix"],
            ["4.0 – 6.9", "Medium", "Fix nel prossimo sprint"],
            ["7.0 – 8.9", "High", "Patcha entro 30 giorni"],
            ["9.0 – 10.0", "CRITICAL", "Patcha SUBITO"],
        ])

    add_h2(doc, "Threat modeling — le 4 domande")
    add_para(doc, "Adam Shostack propone 4 domande:")
    for b in [
        "Cosa stiamo costruendo? → disegna il sistema (DFD)",
        "Cosa può andare storto? → applica STRIDE",
        "Cosa facciamo a riguardo? → mitiga / accetta / elimina",
        "Abbiamo fatto un buon lavoro? → review e iterazione",
    ]:
        add_bullet(doc, b)

    add_callout(doc, "info",
        "Si fa PRIMA di scrivere codice. Su carta o lavagna, in 30-60 minuti. "
        "Trenta minuti che possono evitare ore di refactor.")

    add_h2(doc, "STRIDE — sei lettere, sei categorie")
    add_styled_table(doc,
        ["Lettera", "Categoria", "Esempio"],
        [
            ["S", "Spoofing", "Fingersi qualcun altro (account takeover)"],
            ["T", "Tampering", "Modificare dati (cookie, payload)"],
            ["R", "Repudiation", "Negare di aver fatto un'azione"],
            ["I", "Information Disclosure", "Esporre dati (stack trace, SQLi)"],
            ["D", "Denial of Service", "Rendere indisponibile (DDoS, slowloris)"],
            ["E", "Elevation of Privilege", "IDOR, bypass authz"],
        ])

    add_h2(doc, "STRIDE light — esempio pratico")
    add_para(doc,
        "Prendiamo un sistema piccolo: un blog con login. L'utente si registra, "
        "fa login, scrive post, legge post di altri, lascia commenti.")

    add_image(doc, "img/cap2_dfd_mini_blog.png",
              caption="Data Flow Diagram del blog con trust boundary",
              width_cm=16)

    add_para(doc,
        "Adesso applichiamo STRIDE a ogni elemento del DFD e compiliamo una "
        "tabella di minacce e mitigazioni:")

    add_styled_table(doc,
        ["Elemento", "STRIDE", "Minaccia", "Difesa"],
        [
            ["Utente", "S", "Account takeover via password rubata", "bcrypt + MFA"],
            ["Utente", "R", "'Non ho cancellato io quel post'", "Audit log con IP + ts"],
            ["Webapp", "T", "Modifica cookie di sessione", "Cookie firmato server"],
            ["Webapp", "I", "Stack trace su errore 500", "Error handler generico"],
            ["Webapp", "D", "Brute force login", "Rate limit 5/min"],
            ["Webapp", "E", "SQLi → admin", "Query parametrizzate"],
            ["DB Users", "I", "Backup non cifrato esposto", "Cifratura backup"],
            ["Flusso utente→web", "I", "Sniffing Wi-Fi", "HTTPS (TLS 1.2+)"],
        ])

    add_callout(doc, "tip",
        "In 15 minuti hai prodotto una mappa che vale ore di refactor. "
        "Ricordati: STRIDE non è una formula magica, è un checklist che ti "
        "aiuta a non dimenticare nessuna categoria di minaccia.")

    doc.add_page_break()

    # === CAPITOLO 3 — SQL INJECTION ===
    add_chapter_cover(doc, "3", "SQL Injection",
                       "La vulnerabilità #1 dal 2003: capirla e correggerla",
                       "Lezione 3 · 2 ore")

    add_h2(doc, "Cosa imparerai")
    for b in [
        "Cos'è una SQL Injection e perché è la #1 vulnerabilità del web dal 2003.",
        "Come si esegue un login bypass con ' OR '1'='1' --.",
        "Come si estraggono dati con UNION SELECT.",
        "Perché filtrare gli apici NON funziona.",
        "Come si difende davvero: query parametrizzate.",
    ]:
        add_bullet(doc, b)

    add_h2(doc, "Una storia: il modulo cartaceo")

    add_callout(doc, "story",
        "Immagina di lavorare all'anagrafe. Un cittadino ti porta un modulo "
        "dove ha scritto: 'Mario --- distruggi tutti i moduli precedenti ---'. "
        "Cosa fai? Ovviamente ignori la parte 'in basso': è un modulo, sai "
        "dove finiscono i dati che ti interessano. NON confondi i DATI del "
        "cittadino con le ISTRUZIONI dell'ufficio.")

    add_para(doc,
        "Ora immagina di essere un database. Ricevi una query 'gentile' e "
        "trovi il record. Ma se la query contiene istruzioni iniettate nei "
        "dati, il database meno fortunato le esegue. Questo è SQL Injection.")

    add_h2(doc, "Come si manifesta nel codice")
    add_image(doc, "img/cap3_sqli_vulnerable.png",
              caption="Codice Flask vulnerabile a SQL Injection",
              width_cm=15)

    add_para(doc,
        "Il problema è alla riga della query: stiamo costruendo la stringa "
        "SQL mescolando struttura SQL e input dell'utente con una f-string. "
        "Se l'utente è 'gentile' tutto funziona. Ma l'utente potrebbe non "
        "essere gentile.")

    add_h2(doc, "L'attacco: anatomia in 3 passi")
    add_image(doc, "img/cap3_sqli_anatomia.png",
              caption="Come funziona un attacco SQL Injection di tipo login bypass",
              width_cm=16)

    add_h2(doc, "L'attacco in pratica")
    add_image(doc, "img/cap3_sqli_attack.png",
              caption="Login bypass eseguito con curl: senza conoscere la password",
              width_cm=14)

    add_callout(doc, "danger",
        "Login bypass con ' OR '1'='1' -- è in OWASP Top 10 dal 2003. "
        "Equifax, Heartland, TalkTalk: tutti SQLi. Vent'anni dopo, è ancora "
        "la #1.")

    add_h2(doc, "La correzione: query parametrizzate")
    add_image(doc, "img/cap3_sqli_safe.png",
              caption="Codice corretto con query parametrizzata",
              width_cm=15)

    add_para(doc,
        "I '?' sono PLACEHOLDER. Il driver del database si occupa di compilare "
        "la query SQL prima (struttura fissa) e mandare i dati come VALORI "
        "GIÀ TIPATI, mai come SQL.")

    add_callout(doc, "success",
        "Anche se l'utente scrive ' OR '1'='1' --, il driver lo cerca "
        "LETTERALMENTE come stringa. Non lo trova. Login fallito. La SQLi "
        "diventa IMPOSSIBILE PER DESIGN, non perché filtri qualcosa.")

    add_h2(doc, "Cross-linguaggio: la stessa idea")
    add_styled_table(doc,
        ["Linguaggio", "Pattern corretto"],
        [
            ["Python sqlite3", 'cursor.execute("... = ?", (val,))'],
            ["Python psycopg2", 'cursor.execute("... = %s", (val,))'],
            ["Java JDBC", 'PreparedStatement ps = conn.prepareStatement("... = ?")'],
            ["PHP PDO", '$stmt = $pdo->prepare("... = ?")'],
            ["JS better-sqlite3", 'db.prepare("... = ?").get(val)'],
            ["SQLAlchemy (ORM)", 'User.query.filter_by(email=email).first()'],
        ])

    add_h2(doc, "Perché filtrare NON funziona")

    add_styled_table(doc,
        ["Tentativo di blacklist", "Bypass dell'attaccante"],
        [
            ["replace(\"'\", \"\")", "%27 (URL encoded)"],
            ["Stessa cosa", "\\\\' (escape)"],
            ["Stessa cosa", "Doppi apici \""],
            ["Stessa cosa", "Unicode lookalike ʼ"],
            ["Stessa cosa", "Injection numerica: 1 OR 1=1"],
        ])

    add_callout(doc, "warning",
        "Filtrare caratteri è una STRATEGIA PERDENTE. Gli attaccanti hanno "
        "infiniti modi di bypassarla. Funziona solo finché non arriva qualcuno "
        "bravo. Usa whitelist, meglio ancora parametrizzazione.")

    doc.add_page_break()

    # === CAPITOLO 4 — IDOR + PASSWORD ===
    add_chapter_cover(doc, "4", "Autorizzazione e password",
                       "IDOR, status code, bcrypt — gli errori più costosi",
                       "Lezione 4 · 2 ore")

    add_h2(doc, "Cosa imparerai")
    for b in [
        "La differenza tra autenticazione (sei tu?) e autorizzazione (puoi fare X?).",
        "Cos'è un IDOR e come si corregge.",
        "I codici HTTP 401 vs 403 e quando usarli.",
        "Perché MD5 e SHA-256 NON vanno per password.",
        "Come hashare correttamente con bcrypt.",
    ]:
        add_bullet(doc, b)

    add_h2(doc, "Autenticazione vs Autorizzazione")
    add_para(doc,
        "Sono due cose diverse. Confonderle è il bug concettuale numero 1 dei "
        "programmatori junior.")

    add_styled_table(doc,
        ["Aspetto", "Authentication", "Authorization"],
        [
            ["Domanda", "Chi sei?", "Cosa puoi fare?"],
            ["Quando", "Al login", "A ogni richiesta dopo login"],
            ["Esempio", "Email + password", "user.id == fattura.owner_id"],
            ["Errore tipico", "Login con password rubata", "Vedi dati di altri (IDOR)"],
        ])

    add_h2(doc, "IDOR: Insecure Direct Object Reference")
    add_para(doc,
        "IDOR è il caso classico di authz mancante. L'utente loggato cambia "
        "un identificatore nell'URL e accede a risorse che non sono sue.")

    add_callout(doc, "danger",
        "Caso reale italiano (2022): un e-commerce italiano ha URL "
        "/ordine/<id> non protetti. Cambiando l'ID si vedevano ordini di "
        "altri clienti, con indirizzi, prodotti, importi. Multa Garante "
        "Privacy: ~100.000€. Una sola riga di codice in più — il controllo "
        "di proprietà — avrebbe risparmiato 100.000€.")

    add_h2(doc, "La correzione: ownership check")
    add_image(doc, "img/cap4_idor_safe.png",
              caption="Ownership check con filter_by(owner_id=...)",
              width_cm=15)

    add_h2(doc, "Status code: 401 vs 403 vs 404")
    add_styled_table(doc,
        ["Code", "Significato", "Quando usarlo"],
        [
            ["401 Unauthorized", "Non autenticato", "Manca login o token scaduto"],
            ["403 Forbidden", "Autenticato ma senza permessi", "User normale prova /admin"],
            ["404 Not Found", "Risorsa inesistente", "URL inesistente"],
        ])

    add_h2(doc, "Password hashing: tre operazioni DIVERSE")

    add_styled_table(doc,
        ["", "Encoding", "Encryption", "Hashing"],
        [
            ["Reversibile?", "Sì (banale)", "Sì (con chiave)", "NO"],
            ["Per password?", "❌", "❌", "✅"],
            ["Esempio", "Base64", "AES", "bcrypt"],
            ["Quando", "Trasporto", "Confidenzialità", "Password, integrità"],
        ])

    add_callout(doc, "warning",
        "Base64 NON è cifratura. Chiunque può decodificare in 5 secondi. "
        "Le password vanno HASHATE, mai cifrate.")

    add_h2(doc, "Perché MD5 e SHA-256 NON vanno")
    add_para(doc, "Sentirai dire: 'uso SHA-256, è un hash sicuro'. Per le password è sbagliato:")
    for b in [
        "Velocità: SHA-256 su GPU = miliardi al secondo → brute force in poche ore.",
        "Rainbow tables: senza salt, password identiche → hash identici.",
        "MD5 e SHA-1: collisioni note. Definitivamente morti.",
    ]:
        add_bullet(doc, b)

    add_h2(doc, "bcrypt: la scelta sicura")
    add_image(doc, "img/cap4_bcrypt.png",
              caption="Hashing password con bcrypt: salt automatico, cost configurabile",
              width_cm=15)

    add_callout(doc, "success",
        "bcrypt cost=12 impiega ~250ms per hash. Per l'utente legittimo "
        "(1 hash per login) è impercettibile. Per un attaccante che vuole "
        "fare brute force su milioni di password, è devastante.")

    add_h2(doc, "Diagnosi visiva nel DB")
    add_styled_table(doc,
        ["Cosa vedi", "Diagnosi"],
        [
            ["mariopwd (testo leggibile)", "🔥 In chiaro — catastrofico"],
            ["5f4dcc3b... (32 hex)", "🔥 MD5 — morto"],
            ["5baa61e4... (40 hex)", "🔥 SHA-1 — morto"],
            ["e3b0c44... (64 hex)", "⚠ SHA-256 senza salt — inadeguato"],
            ["$2b$12$KIXbN...", "✅ bcrypt — OK"],
            ["$argon2id$v=19$...", "✅ Argon2id — ottimo"],
        ])

    doc.add_page_break()

    # === CAPITOLO 5 — XSS + HEADER ===
    add_chapter_cover(doc, "5", "XSS e header HTTP di sicurezza",
                       "JavaScript dell'attaccante nel browser della vittima",
                       "Lezione 5 · 2 ore")

    add_h2(doc, "Cosa imparerai")
    for b in [
        "Cos'è il Cross-Site Scripting e i suoi 3 tipi.",
        "Come si esegue una XSS Reflected e una XSS Stored.",
        "Come si difende con escape automatico dei template engine.",
        "I sei header HTTP di sicurezza principali.",
        "Come configurare cookie sicuri (Secure, HttpOnly, SameSite).",
    ]:
        add_bullet(doc, b)

    add_h2(doc, "Cos'è XSS")
    add_para(doc,
        "XSS è quando l'attaccante riesce a far eseguire del SUO JavaScript "
        "dentro la pagina del TUO sito. Il browser della vittima esegue il "
        "codice come se fosse fidato, dando all'attaccante accesso a cookie, "
        "DOM, possibilità di fare richieste a nome dell'utente.")

    add_h2(doc, "I tre tipi di XSS")
    add_styled_table(doc,
        ["Tipo", "Dove sta il payload", "Vittima"],
        [
            ["Reflected", "Nell'URL", "Chi clicca un link malevolo"],
            ["Stored", "Nel database", "Chiunque visita la pagina (più grave)"],
            ["DOM-based", "Manipolazione client-side", "Più raro, più subdolo"],
        ])

    add_h2(doc, "XSS Stored: il flusso dell'attacco")
    add_image(doc, "img/cap5_xss_flow.png",
              caption="Anatomia di un attacco XSS Stored e le sue difese",
              width_cm=16)

    add_h2(doc, "La difesa primaria: escape dell'output")
    add_para(doc,
        "I template engine moderni fanno escape AUTOMATICO. In Flask, il "
        "template engine si chiama Jinja2:")

    add_code_block(doc,
        """<!-- Template Jinja2 -->
<h1>Risultati per: {{ query }}</h1>

# Se query = "<script>alert(1)</script>"
# Jinja2 lo trasforma in: &lt;script&gt;alert(1)&lt;/script&gt;
# Il browser lo mostra come TESTO, non lo esegue""")

    add_callout(doc, "danger",
        "Il filtro |safe in Jinja2 (e equivalenti come v-html in Vue, "
        "dangerouslySetInnerHTML in React) DISABILITA l'escape. Usato su "
        "input utente = XSS GARANTITA. Mai farlo.")

    add_h2(doc, "I 6 header HTTP di sicurezza")
    add_styled_table(doc,
        ["Header", "Cosa fa"],
        [
            ["Strict-Transport-Security", "Forza HTTPS dopo prima visita (HSTS)"],
            ["Content-Security-Policy", "Limita risorse caricabili (anti-XSS L2)"],
            ["X-Frame-Options: DENY", "Blocca clickjacking via iframe"],
            ["X-Content-Type-Options: nosniff", "Blocca MIME sniffing"],
            ["Referrer-Policy", "Controlla cosa va nel Referer"],
            ["Permissions-Policy", "Limita API browser (camera, mic, geoloc)"],
        ])

    add_h2(doc, "Cookie sicuri: i 3 attributi")
    add_code_block(doc,
        """# Configurazione Flask
app.config.update(
    SESSION_COOKIE_SECURE=True,      # Solo HTTPS
    SESSION_COOKIE_HTTPONLY=True,     # JS NON può leggerlo (anti-XSS)
    SESSION_COOKIE_SAMESITE="Lax",    # Anti-CSRF
)""")

    add_callout(doc, "tip",
        "Test rapido del tuo sito: vai su securityheaders.com e inserisci il "
        "dominio. Ti dà un voto da F ad A+. Github prende A, molti siti "
        "italiani PA prendono F. Aggiungere tutti gli header in Flask: "
        "10 minuti con Flask-Talisman.")

    doc.add_page_break()

    # === CAPITOLO 6 — VALIDATION + SUPPLY CHAIN ===
    add_chapter_cover(doc, "6", "Input validation e supply chain",
                       "Pydantic, path traversal, il rischio delle dipendenze",
                       "Lezione 6 · 2 ore")

    add_h2(doc, "Cosa imparerai")
    for b in [
        "Distinzione tra validation, sanitization e encoding.",
        "Perché whitelist è meglio di blacklist.",
        "Usare Pydantic per validation strutturata in Python.",
        "Cos'è il Path Traversal e come si corregge.",
        "Scoprire CVE nelle dipendenze con pip-audit.",
    ]:
        add_bullet(doc, b)

    add_h2(doc, "Validation, sanitization, encoding")
    add_styled_table(doc,
        ["Operazione", "Cosa fa", "Quando si usa"],
        [
            ["Validation", "Verifica + rifiuta input non validi", "All'ENTRATA"],
            ["Sanitization", "Modifica per renderlo sicuro", "Solo se serve mantenere struttura ricca"],
            ["Encoding", "Trasforma in output a seconda del contesto", "All'USCITA"],
        ])

    add_callout(doc, "tip",
        "Regola d'oro: VALIDATE all'entrata, ENCODE all'uscita, SANITIZE solo "
        "se l'input deve mantenere struttura ricca (HTML, markdown).")

    add_h2(doc, "Pydantic: validation strutturata")
    add_image(doc, "img/cap6_pydantic.png",
              caption="Modello Pydantic con validator custom",
              width_cm=15)

    add_para(doc, "Cosa fa Pydantic per te:")
    for b in [
        "Controlla i TIPI (se non corretti, errore)",
        "Controlla i VINCOLI (lunghezza, range, regex)",
        "Esegue VALIDATOR custom (regole business)",
        "Restituisce errori dettagliati con path JSON",
        "Genera automaticamente JSON Schema (per OpenAPI/FastAPI)",
    ]:
        add_bullet(doc, b)

    add_h2(doc, "Path Traversal e correzione")
    add_para(doc,
        "Path traversal è una vulnerabilità negli endpoint che servono file "
        "dall'input dell'utente. Codice vulnerabile:")

    add_code_block(doc,
        """@app.route("/download")
def download():
    filename = request.args.get("file")
    return send_file(f"./uploads/{filename}")

# Attacco: GET /download?file=../etc/passwd
# Path costruito: ./uploads/../etc/passwd
# Risolto in: /etc/passwd → l'attaccante legge file di sistema""")

    add_para(doc, "La correzione richiede tre controlli:")
    add_code_block(doc,
        """import os
UPLOAD_DIR = os.path.realpath("./uploads")
ALLOWED_EXTS = {".pdf", ".png", ".jpg"}

@app.route("/download")
@login_required
def download():
    filename = request.args.get("file", "")
    # 1) Whitelist: niente separatori, no nascosti
    if "/" in filename or "\\\\" in filename or filename.startswith("."):
        abort(400)
    # 2) Whitelist estensione
    if os.path.splitext(filename)[1].lower() not in ALLOWED_EXTS:
        abort(400)
    # 3) realpath + startswith verifica che resti dentro UPLOAD_DIR
    full = os.path.realpath(os.path.join(UPLOAD_DIR, filename))
    if not full.startswith(UPLOAD_DIR + os.sep):
        abort(403)
    return send_from_directory(UPLOAD_DIR, filename, as_attachment=True)""")

    add_h2(doc, "Supply chain: il rischio nascosto")
    add_para(doc,
        "Un progetto Python medio ha 50-200 dipendenze indirette. Ogni dipendenza "
        "è codice scritto da qualcun altro. Se ha una CVE, la tua app è vulnerabile.")

    add_callout(doc, "story",
        "Dicembre 2021: viene pubblicata CVE-2021-44228 in Log4j (Java). "
        "CVSS 10.0. Una stringa in un header HTTP permetteva di eseguire "
        "codice arbitrario sul server. Per due settimane, mezza Internet ha "
        "patchato in emergenza.")

    add_h2(doc, "pip-audit in azione")
    add_image(doc, "img/cap6_pip_audit.png",
              caption="Scansione delle dipendenze con pip-audit",
              width_cm=14)

    add_callout(doc, "info",
        "Dal dicembre 2027 il Cyber Resilience Act renderà l'SBOM obbligatorio "
        "per ogni prodotto digitale venduto in UE. Iniziate a familiarizzare con "
        "pip-audit e cyclonedx-bom già adesso.")

    doc.add_page_break()

    # === CAPITOLO 7 — DOCUMENTAZIONE + AI ===
    add_chapter_cover(doc, "7", "Documentazione di sicurezza e uso dell'IA",
                       "SECURITY.md, validare codice generato da AI",
                       "Lezione 7 · 2 ore")

    add_h2(doc, "Cosa imparerai")
    for b in [
        "Perché documentare i controlli di sicurezza è un requisito.",
        "Come strutturare un documento SECURITY.md.",
        "Come l'IA può aiutare a scrivere codice (e dove può fregarti).",
        "I sette errori tipici del codice generato da IA.",
        "Quando NON usare l'IA.",
    ]:
        add_bullet(doc, b)

    add_h2(doc, "Documentare la sicurezza")
    add_para(doc,
        "Nei sei capitoli precedenti hai imparato decine di pattern. Adesso: "
        "come fai a dimostrare che la tua app li ha tutti?")

    add_callout(doc, "info",
        "Non è burocrazia, è un REQUISITO LEGALE: GDPR Art. 32, NIS 2 Art. 21, "
        "Cyber Resilience Act 2027. E un REQUISITO OPERATIVO: quando arriva un "
        "auditor, deve poter verificare in mezza giornata.")

    add_h2(doc, "Template SECURITY.md — 9 sezioni")
    add_styled_table(doc,
        ["#", "Sezione"],
        [
            ["1", "Informazioni generali (progetto, owner, dati trattati)"],
            ["2", "Threat model (DFD + tabella STRIDE)"],
            ["3", "Controlli applicati (auth, authz, validation, ecc.)"],
            ["4", "Vulnerabilità note e debiti tecnici"],
            ["5", "Test di sicurezza (automatici + pentest)"],
            ["6", "Incident response (contatti, playbook)"],
            ["7", "Compliance (GDPR, NIS 2, CRA)"],
            ["8", "Approvazione e revisione"],
            ["9", "Allegati (DPIA, pentest report, SBOM)"],
        ])

    add_callout(doc, "tip",
        "Esercizio raccomandato: per il TUO progetto di stage, mantieni un "
        "SECURITY.md compilato fin dall'inizio. Sarà uno degli output più "
        "apprezzati in azienda — fa una differenza enorme rispetto a chi "
        "consegna solo codice 'che funziona'. Collegamento con UF 7 "
        "(Tecniche di redazione documentazione tecnica).")

    add_h2(doc, "Uso responsabile dell'IA")
    add_para(doc,
        "Nel 2026 scrivere codice senza assistenti AI è raro. Ma usare AI "
        "senza saperla validare è pericoloso. L'IA è addestrata su miliardi "
        "di righe di codice pubblico, INCLUSO codice vulnerabile.")

    add_callout(doc, "warning",
        "Studi GitHub (2022), Stanford (2023): ~40% dei suggerimenti Copilot "
        "in scenari di sicurezza contengono vulnerabilità. Gli sviluppatori "
        "con AI scrivono codice leggermente MENO sicuro, ma sono PIÙ "
        "convinti che sia sicuro (bias cognitivo).")

    add_h2(doc, "I 7 errori tipici del codice IA")
    add_styled_table(doc,
        ["#", "Errore", "Soluzione"],
        [
            ["1", "SQL Injection con f-string", "Query parametrizzate"],
            ["2", "Hash deboli (SHA-256) per password", "bcrypt / Argon2id"],
            ["3", "Manca authorization check", "Ownership check"],
            ["4", "Template senza escape", "Jinja2/Twig auto-escape"],
            ["5", "CORS troppo permissivo ('*')", "Whitelist origini"],
            ["6", "Catch-all che falliscono in modo aperto", "Fail Secure"],
            ["7", "Segreti hardcoded ('change-me')", "Env var"],
        ])

    add_h2(doc, "Workflow di validazione in 4 step")
    for b in [
        "Step 1 — LEGGI E CAPISCI (30 sec). Se non sapresti spiegarlo, non accettare.",
        "Step 2 — SCANSIONA PATTERN (1 min). f-string in SQL? Hash deboli? eval?",
        "Step 3 — VERIFICA CONTESTO (1-2 min). Coerente con architettura, librerie?",
        "Step 4 — TEST + LINTER (5-15 min). pytest + bandit/semgrep.",
    ]:
        add_bullet(doc, b)

    add_callout(doc, "danger",
        "Quando NON usare l'IA: crittografia 'fatta in casa', codice di sicurezza "
        "critico, compliance/legale. E quando NON sai validare: impara prima!")

    doc.add_page_break()

    # === CAPITOLO 8 — LAB INTEGRATO ===
    add_chapter_cover(doc, "8", "Lab integrato e verifica finale",
                       "Sei un junior security analyst",
                       "Lezione 8 · 2 ore — VERIFICA")

    add_h2(doc, "Lo scenario")
    add_para(doc,
        "Hai finito le sette lezioni. Sai cosa è una SQL Injection, sai "
        "correggerla; sai cosa è un IDOR, sai correggerlo; sai cosa è un XSS, "
        "sai difenderti. In questa lezione APPLICHI tutto, su un'app vera, in "
        "80 minuti.")

    add_callout(doc, "info",
        "Sei un junior security analyst. Ti viene affidata una piccola app "
        "per una review di sicurezza prima del rilascio. Output: un mini-report "
        "scritto (1-2 pagine) con almeno 3 vulnerabilità identificate, PoC, "
        "fix proposto.")

    add_h2(doc, "Le regole del lab")
    add_styled_table(doc,
        ["✅ Permesso", "🚫 Non permesso"],
        [
            ["Leggere il codice sorgente", "Aprire la versione corretta"],
            ["DevTools, curl, sqlite3, DB Browser", "Cercare la soluzione su Google"],
            ["Lavorare a coppie", "Copiare il report"],
            ["1 hint gratis dal docente", "2° hint costa -5% sul voto"],
            ["Consultare le dispense", ""],
        ])

    add_h2(doc, "Griglia di valutazione")
    add_styled_table(doc,
        ["Voce", "Peso"],
        [
            ["Numero vulnerabilità (≥3 = sufficiente, 5+ = ottimo)", "25"],
            ["Correttezza tecnica dei PoC", "30"],
            ["Qualità dei fix proposti", "30"],
            ["Severity giustificata coerentemente", "10"],
            ["Mapping OWASP / norma violata", "5"],
        ])
    add_para(doc, "Soglia di sufficienza: 60/100.", bold=True)

    add_h2(doc, "Cosa porti via dal corso")
    add_para(doc, "Cinque idee, in cinque anni ti resteranno queste:")
    for b in [
        "1. La sicurezza si PROGETTA dall'inizio, non si aggiunge alla fine.",
        "2. DEFENSE IN DEPTH sempre. Mai una sola difesa.",
        "3. MENTALITÀ AVVERSARIA. Cosa fa che non dovrebbe?",
        "4. I fondamentali OWASP — riconosci a vista, correggi a memoria.",
        "5. DOCUMENTA, TESTA, AUTOMATIZZA.",
    ]:
        add_bullet(doc, b)

    add_callout(doc, "tip",
        "Per crescere ancora (gratuito): PortSwigger Web Security Academy "
        "(il miglior corso al mondo), TryHackMe, HackTheBox starting point, "
        "OWASP cheat sheets. Certificazioni entry: CompTIA Security+, eJPT, "
        "PortSwigger BSCP.")

    doc.add_page_break()

    # === APPENDICI ===
    add_h1(doc, "Appendice A — Snippet di pronto utilizzo")

    add_h2(doc, "Setup ambiente in 5 comandi")
    add_code_block(doc,
        """mkdir progetto && cd progetto
python -m venv .venv
.\\.venv\\Scripts\\Activate.ps1     # Windows
# source .venv/bin/activate        # macOS/Linux
pip install flask pydantic[email] bcrypt bleach pip-audit""")

    add_h2(doc, "Query parametrizzata")
    add_code_block(doc, """cursor.execute("SELECT * FROM users WHERE id = ?", (user_id,))""")

    add_h2(doc, "Ownership check")
    add_code_block(doc,
        """f = Fattura.query.filter_by(
    id=fid, owner_id=session["user_id"]
).first_or_404()""")

    add_h2(doc, "bcrypt")
    add_code_block(doc,
        """import bcrypt
h = bcrypt.hashpw(pwd.encode(), bcrypt.gensalt(rounds=12))
ok = bcrypt.checkpw(input_pwd.encode(), h)""")

    add_h2(doc, "Cookie sicuri Flask")
    add_code_block(doc,
        """app.config.update(
    SESSION_COOKIE_SECURE=True,
    SESSION_COOKIE_HTTPONLY=True,
    SESSION_COOKIE_SAMESITE="Lax",
)""")

    doc.add_page_break()

    add_h1(doc, "Appendice B — Riferimenti normativi")
    add_styled_table(doc,
        ["Articolo", "Cosa richiede"],
        [
            ["GDPR Art. 5", "Principi: minimizzazione, integrità, riservatezza"],
            ["GDPR Art. 25", "Privacy by Design and by Default"],
            ["GDPR Art. 32", "Misure tecniche adeguate (cifratura, pseudonimizzazione)"],
            ["GDPR Art. 33-34", "Notifica breach al Garante entro 72h"],
            ["NIS 2 Art. 21", "10 categorie di misure di gestione del rischio"],
            ["NIS 2 Art. 23", "Notifica incidenti: 24h + 72h + 30 giorni"],
            ["CRA (dic. 2027)", "Niente vulnerabilità note alla vendita, SBOM obbligatorio"],
            ["L. 4/2004", "Accessibilità prodotti digitali (PA)"],
        ])

    add_h1(doc, "Appendice C — Glossario essenziale")
    add_styled_table(doc,
        ["Termine", "Significato"],
        [
            ["CIA Triad", "Confidentiality, Integrity, Availability"],
            ["CVE", "Common Vulnerabilities and Exposures — identificatore vuln"],
            ["CVSS", "Common Vulnerability Scoring System — punteggio 0-10"],
            ["DFD", "Data Flow Diagram"],
            ["IDOR", "Insecure Direct Object Reference"],
            ["OWASP", "Open Web Application Security Project"],
            ["RCE", "Remote Code Execution"],
            ["SAST", "Static Application Security Testing"],
            ["DAST", "Dynamic Application Security Testing"],
            ["SBOM", "Software Bill of Materials"],
            ["STRIDE", "Spoofing/Tampering/Repudiation/Info/DoS/Elevation"],
            ["XSS", "Cross-Site Scripting"],
            ["SQLi", "SQL Injection"],
            ["MFA", "Multi-Factor Authentication"],
            ["HSTS", "HTTP Strict Transport Security"],
            ["CSP", "Content-Security-Policy"],
        ])

    add_h1(doc, "Appendice D — Risorse esterne")
    add_h2(doc, "Documentazione")
    for b in [
        "OWASP Top 10 — https://owasp.org/Top10",
        "OWASP Cheat Sheet Series — https://cheatsheetseries.owasp.org",
        "PortSwigger Web Security Academy — https://portswigger.net/web-security",
        "TryHackMe — https://tryhackme.com",
        "Garante Privacy — https://www.garanteprivacy.it",
    ]:
        add_bullet(doc, b)

    add_h2(doc, "Repository del corso")
    for b in [
        "https://github.com/ss4i/corso-stem-ifts-secure-coding",
        "https://github.com/ss4i/corso-its-cybersecurity-32h (estensione)",
    ]:
        add_bullet(doc, b)

    # Footer finale
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("━" * 30)
    run.font.color.rgb = hex_to_rgb(COL_ACCENT)
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Fine della dispensa")
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = hex_to_rgb(COL_PRIMARY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Buona strada.")
    run.font.size = Pt(14)
    run.font.italic = True
    run.font.color.rgb = hex_to_rgb(COL_GREY)

    doc.save("dispensa_completa_v2.docx")
    print("OK dispensa_completa_v2.docx")


if __name__ == "__main__":
    os.chdir(os.path.dirname(os.path.abspath(__file__)))
    build_dispensa()
